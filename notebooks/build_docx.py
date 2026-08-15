"""Builds entregables/resumen_ejecutivo.docx from the English executive summary.

Same pattern as case 1: the content is written here rather than parsed from the Markdown,
so the two stay in sync deliberately instead of by regex accident. Colours follow this
case's theme (green) rather than case 1's blue.

Run: python notebooks/build_docx.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

RUTA_BASE = Path(__file__).resolve().parents[1]
GRAFICOS = RUTA_BASE / 'salidas' / 'graficos'
OUT = RUTA_BASE / 'entregables' / 'resumen_ejecutivo.docx'

VERDE = RGBColor(0x0F, 0x7D, 0x3F)      # el acento del caso
VERDE_OSCURO = RGBColor(0x0A, 0x4F, 0x28)
DARKTEXT = RGBColor(0x1A, 0x18, 0x15)
GRAY = RGBColor(0x6E, 0x7B, 0x77)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.color.rgb = DARKTEXT


def heading(text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = VERDE_OSCURO
        run.font.name = 'Cambria'
    return h


def para(text, *, italic=False, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or DARKTEXT
    return p


def bullet(texto_fuerte, resto):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(texto_fuerte)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DARKTEXT
    r2 = p.add_run(resto)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARKTEXT
    return p


def image(filename, width_in=6.0, caption=None):
    doc.add_picture(str(GRAFICOS / filename), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = GRAY


# Title block ---------------------------------------------------------------
title = doc.add_paragraph()
r = title.add_run('The football transfer market: who pays whom, and what age costs')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = VERDE_OSCURO
r.font.name = 'Cambria'

sub = doc.add_paragraph()
r = sub.add_run('Executive summary for the investment committee  ·  15 August 2026')
r.font.size = Pt(12)
r.font.color.rgb = GRAY

doc.add_paragraph()

# What we set out to test ---------------------------------------------------
heading('What we set out to test')
para('Two things everyone in football takes for granted: that transfer money is concentrating in '
     'fewer and fewer clubs, and that young players now carry a growing price premium. Both matter '
     'to this fund, because between them they decide whether "develop and sell" still has a margin '
     'or whether capital belongs at the buying end.')
para('We wrote that expectation down before opening the data. The analysis contradicted it on both '
     'counts.', bold=True)

# Findings ------------------------------------------------------------------
heading('What we found')

para('The money is spreading out, not concentrating.', bold=True)
para('The ten largest buying clubs took 32.1% of European transfer spending in 2022/23 and 26.7% '
     'in 2025/26, while the number of clubs buying at all rose from 339 to 382. This is not an '
     'English effect — excluding English buyers the share still falls, 30.6% to 24.0% — and it is '
     'not a four-season blip: FIFA’s census shows 45% more clubs paying transfer fees in 2025 '
     'than in 2018.')
image('01_concentracion_cae.png',
      caption='The top 10% of buyers took 65% of the money in 2022/23 and 60% in 2025/26.')
image('02_mas_clubes_cada_ano.png',
      caption='FIFA TMS census: clubs paying a fee up 45% since 2018, clubs receiving one up 49%.')

para('Selling is a durable position, not a bad year.', bold=True)
para('Of 350 clubs present in all four seasons, 202 were net sellers in three or more, and 96 in '
     'all four; chance alone would have produced 61. Ajax banked €267m net across the period, '
     'Salzburg €234m, Lille €185m. Chelsea (−€800m) and Manchester United '
     '(−€672m) never once finished a season as net sellers.')
image('03_vender_es_estructural.png',
      caption='Clubs above or below what random role assignment would produce.')

para('Youth carries no premium.', bold=True)
para('An 18–23 player costs roughly what a 24–29 player costs — a ratio around 1.0 that '
     'does not escalate in either of our two independent sources. Clubs paying more for youth as '
     'such are not following the market.')
image('04_prima_juvenil_inexistente.png',
      caption='Ratio of average fees, 18-23 against 24-29, in both sources.')

para('But the young market runs at two speeds.', bold=True)
para('Deals of €40m or more are 2.3% of all young-player transactions and take 23.8% of the '
     'money spent on them, while the 71% priced under €5m account for just 15%. The headline '
     'signings are real; they simply sit on top of a long, cheap tail that is what actually moves '
     'the totals.')
image('05_dos_velocidades.png',
      caption='Share of deals against share of spending, players aged 18-23.')

# Recommendations -----------------------------------------------------------
heading('What we recommend')
bullet('Put the next tranche into a persistent net seller in a mid-tier league, ',
       'not into an elite club. The scarce asset is not money to spend — that side of the market '
       'has more participants every year — but a reliable supply of sellable players.')
bullet('Cap prices against the market band and write into the mandate that age carries no '
       'premium. ', 'A policy change, not a project.')
bullet('Make each club pick its lane ',
       'in the young-player market: cheap volume or elite few. A club drifting between the two '
       'competes badly in both.')

# Limitations ---------------------------------------------------------------
heading('What this analysis cannot tell you')
para('Four seasons of deal-level data is not a trend, so every statement about direction over time '
     'rests on FIFA’s census rather than on ours. We make no claim about the share of spending '
     'going to young players in level or direction: our two sources disagree, and the gap doubles '
     'in 2025 for reasons we cannot defend. Nothing here is causal — the case describes how the '
     'market is structured, not what moves it. And net transfer income is not profit: without club '
     'accounts, "sells well" is not the same as "makes money".')
para('Full method, checks and decision log: CASO.md in the case repository.',
     italic=True, size=10, color=GRAY)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f'escrito {OUT.relative_to(RUTA_BASE)}  ({OUT.stat().st_size / 1024:.0f} KB)')
